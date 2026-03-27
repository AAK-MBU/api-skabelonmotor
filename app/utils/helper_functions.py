"""
Utility helpers used by the Skabelonmotor API.

This module contains functions responsible for:
1. Normalizing lightweight HTML formatting used in the template engine.
2. Rendering the final letter into PDF or DOCX formats.
3. Normalizing keys for reliable comparisons.
4. Replacing placeholders inside generated letter text.
"""

import base64
import io
import re
import tempfile

from io import BytesIO

from bs4 import BeautifulSoup

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import RGBColor, Pt

from docx2pdf import convert

from openpyxl import load_workbook
from openpyxl.cell.rich_text import CellRichText

# Regex used to detect block headers such as:
# "Blok 1", "Blok 3.1", "Blok 7.2a"
BLOCK_HEADER_PATTERN = re.compile(r"^Blok\s+([0-9]+(?:\.\s*[0-9]+)?[a-zA-Z]?)")


def add_hyperlink(paragraph, url, text):

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")

    # Blue color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    # Underline
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def extract_cell_formatting(cell):
    """
    Convert Excel rich text content into HTML-like formatted text.

    Handles formatting such as bold, italic, underline, strike-through
    and color while preserving the original text structure.

    Args:
        cell: openpyxl cell object.

    Returns:
        str: HTML-like formatted text.
    """

    if cell is None or cell.value is None:
        return ""

    value = cell.value

    # ----------------------------------------
    # Rich formatted text (Excel rich text)
    # ----------------------------------------
    if isinstance(value, CellRichText):

        parts = []

        for block in value:

            text = block.text or ""
            font = block.font

            if not text:
                continue

            # Remove zero-width characters sometimes inserted by Excel
            text = text.replace("\u200b", "")

            # Replace Excel tab indentation
            text = text.replace("\t", " ")

            prefix = ""
            suffix = ""

            if font:

                # Bold
                if font.b:
                    prefix += "<strong>"
                    suffix = "</strong>" + suffix

                # Italic
                if font.i:
                    prefix += "<em>"
                    suffix = "</em>" + suffix

                # Underline
                if font.u in ["single", "double", "singleAccounting", "doubleAccounting", True]:
                    prefix += "<u>"
                    suffix = "</u>" + suffix

                # Strikethrough
                if font.strike:
                    prefix += "<strike>"
                    suffix = "</strike>" + suffix

                # Text color
                if font.color and font.color.rgb:
                    rgb = font.color.rgb[-6:]

                    # Skip default black text
                    if rgb != "000000":
                        prefix += f'<span style="color:#{rgb}">'
                        suffix = "</span>" + suffix

            parts.append(f"{prefix}{text}{suffix}")

        return "".join(parts)

    # ----------------------------------------
    # Plain text cell (no formatting)
    # ----------------------------------------
    return str(value)


def parse_workbook_afgoerelsesbrev(binary_excel: bytes) -> list[dict]:
    """
    Pure Excel parser.

    Extracts blocks and their entries from the workbook without applying
    any business logic, metadata, or custom functions.

    Args:
        binary_excel (bytes): Excel workbook content.

    Returns:
        list[dict]: Raw extracted block structures.
    """

    LINK_MAPPING = {
        "Folkeskoleloven (retsinformation.dk)": "https://www.retsinformation.dk/eli/lta/2025/1100#P26",

        "Bekendtgørelse om befordring af elever i folkeskolen (retsinformation.dk)": "https://www.retsinformation.dk/eli/lta/2014/688",

        "Ungdomsskoleloven (retsinformation.dk)": "https://www.retsinformation.dk/eli/lta/2010/665",

        "Behandling af personoplysninger i Børn og Unge (aarhus.dk)": "https://aarhus.dk/om-kommunen/databeskyttelse/behandling-af-personoplysninger-i-boern-og-unge"
    }

    def inject_links(entry_text: str) -> str:
        for text, url in LINK_MAPPING.items():
            if text in entry_text:
                print(f"text: {text}")
                print(f"url: {url}")
                print(f"entry_text:\n{entry_text}")
                print()

                entry_text = entry_text.replace(
                    text,
                    f'<a href="{url}">{text}</a>'
                )
                print(f"after:\n{entry_text}")
                print()
                print()
                print()

        return entry_text

    wb = load_workbook(BytesIO(binary_excel), rich_text=True)

    parsed_blocks = []
    current_block = None

    # ----------------------------------------
    # Parse workbook sheets
    # ----------------------------------------
    for sheet_name in wb.sheetnames:

        if not sheet_name.startswith("Blok"):
            continue

        ws = wb[sheet_name]
        rows = list(ws.iter_rows())

        for i, row in enumerate(rows):

            col_a_cell = row[0] if len(row) > 0 else None
            col_b_cell = row[1] if len(row) > 1 else None

            col_a = col_a_cell.value if col_a_cell else None
            col_b = extract_cell_formatting(col_b_cell) if col_b_cell else None

            # ----------------------------------------
            # Detect block header
            # ----------------------------------------
            if isinstance(col_a, str):

                match = BLOCK_HEADER_PATTERN.match(col_a)

                if match:

                    block_id = match.group(1).replace(" ", "").strip()

                    # Mapping key from column C in next row
                    next_row = rows[i + 1] if i + 1 < len(rows) else None
                    next_col_c = None

                    if next_row and len(next_row) > 2:
                        next_col_c = next_row[2].value

                    current_block = {
                        "block_id": block_id,
                        "title": col_a,
                        "mapping": str(next_col_c).strip() if next_col_c else None,
                        "entries": {}
                    }

                    parsed_blocks.append(current_block)

                    continue

            if not current_block:
                continue

            # ----------------------------------------
            # Parse entries
            # ----------------------------------------
            if col_a and col_b:

                entry_text = col_b.strip()

                # Skip "Ingen tekst"
                if normalize_key(entry_text) == "ingentekst":
                    continue

                key = str(col_a)

                entry_text = inject_links(entry_text)

                current_block["entries"][key] = entry_text

    return parsed_blocks


def insert_letter_into_template(template_b64: str, letter_text: str) -> bytes:
    """
    Insert rendered HTML letter text into a DOCX template.

    The function locates the {{LETTER_TEXT}} placeholder in the template,
    removes it, and inserts the formatted HTML content at the same location.
    All supported formatting (bold, italic, underline, strike, color)
    is preserved using the same recursive HTML parsing logic used by the
    standalone DOCX renderer.
    """

    template_bytes = base64.b64decode(template_b64)

    doc = Document(BytesIO(template_bytes))

    # -------------------------------------------------
    # Recursive HTML → DOCX run processor
    # -------------------------------------------------
    def process_node(node, paragraph, formatting=None):

        if formatting is None:
            formatting = {}

        # ------------------------------
        # TEXT NODE
        # ------------------------------
        if node.name is None:

            content = str(node)

            if not content.strip():
                return

            run = paragraph.add_run(content)

            if formatting.get("bold"):
                run.bold = True

            if formatting.get("italic"):
                run.italic = True

            if formatting.get("underline"):
                run.underline = True

            if formatting.get("strike"):
                run.font.strike = True

            rgb = formatting.get("color")

            if isinstance(rgb, str) and len(rgb) == 6:
                run.font.color.rgb = RGBColor.from_string(rgb)

        # ------------------------------
        # ELEMENT NODE
        # ------------------------------
        else:

            # ----------------------------------------
            # ELEMENT NODE
            # ----------------------------------------
            # When encountering an HTML element (<b>, <span>, etc.)
            # we update the active formatting before processing children.
            # Child nodes inherit this updated formatting.
            new_format = formatting.copy()

            if node.name == "a":

                url = node.get("href")

                # Create empty hyperlink first
                hyperlink_text = ""

                for child in node.children:
                    hyperlink_text += str(child)

                add_hyperlink(paragraph, url, hyperlink_text)

                return

            if node.name in ["strong", "b"]:
                new_format["bold"] = True

            if node.name in ["em", "i"]:
                new_format["italic"] = True

            if node.name == "u":
                new_format["underline"] = True

            if node.name == "strike":
                new_format["strike"] = True

            if node.name in ["span", "font"]:
                match = re.search(r"#([0-9A-Fa-f]{6})", str(node))
                if match:
                    new_format["color"] = match.group(1)

            # Recursively process child nodes so formatting cascades
            for child in node.children:
                process_node(child, paragraph, new_format)

    # -------------------------------------------------
    # Find placeholder and insert content
    # -------------------------------------------------
    for paragraph in doc.paragraphs:

        if "{{LETTER_TEXT}}" in paragraph.text.upper():

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            # Remove placeholder paragraph
            parent.remove(paragraph._element)

            paragraphs = letter_text.split("\n\n")

            for offset, p in enumerate(paragraphs):

                new_paragraph = doc.add_paragraph()

                soup = BeautifulSoup(p, "html.parser")

                for child in soup.children:
                    process_node(child, new_paragraph)

                # Move paragraph to correct location
                parent.insert(index + offset, new_paragraph._element)

                new_paragraph.paragraph_format.space_after = Pt(12)  # or 6, 18 etc

            break

    buffer = BytesIO()
    doc.save(buffer)

    return buffer.getvalue()


def normalize_html(text: str) -> str:
    """
    Normalize lightweight HTML formatting so it works with the rendering engines.

    The template system produces simple HTML-like formatting (span, strong, em),
    but ReportLab and python-docx support slightly different tag sets. This
    function converts unsupported tags into equivalents that both renderers
    understand.

    Args:
        text (str): HTML-like formatted text.

    Returns:
        str: Normalized HTML ready for rendering.
    """

    # ReportLab does not support <span style="color:..."> but does support <font>
    text = re.sub(
        r'<span style="color:#([0-9A-Fa-f]{6})">',
        r'<font color="#\1">',
        text
    )

    # Close converted color tags
    text = text.replace("</span>", "</font>")

    # ReportLab expects <b> and <i> rather than <strong> / <em>
    text = text.replace("<strong>", "<b>").replace("</strong>", "</b>")
    text = text.replace("<em>", "<i>").replace("</em>", "</i>")

    return text


def html_to_docx_bytes(text: str) -> bytes:
    """
    Render HTML-like formatted text into a DOCX document.

    The function converts the lightweight HTML produced by the template
    system into a Word document using python-docx. The text is first split
    into paragraphs, after which each paragraph is parsed as HTML and
    traversed node-by-node.

    A nested helper function (`process_node`) is used to recursively walk
    the HTML tree and translate formatting tags into python-docx run
    formatting (bold, italic, underline, strike, color). This recursive
    traversal ensures that nested formatting is preserved correctly.

    Args:
        text (str): HTML-like formatted letter text.

    Returns:
        bytes: Generated DOCX file content.
    """

    soup = BeautifulSoup(text, "html.parser")

    doc = Document()

    def process_node(node, paragraph, formatting=None):
        """
        Recursively process HTML nodes and convert them into DOCX runs.

        This function walks the HTML DOM tree produced by BeautifulSoup.
        As it descends through the tree, it keeps track of the currently
        active formatting (bold, italic, color, etc.) so nested formatting
        is preserved when generating DOCX runs.

        The recursion works like this:

            HTML structure
                ↓
            process_node(node)
                ↓
            update formatting if node is a tag
                ↓
            call process_node(child) for each child

        Args:
            node: BeautifulSoup node (text or HTML element).
            paragraph: python-docx paragraph where runs are inserted.
            formatting (dict | None): Active formatting inherited from
                parent nodes.
        """

        if formatting is None:
            formatting = {}

        # ----------------------------------------
        # TEXT NODE
        # ----------------------------------------
        # If the node has no tag name, it represents plain text
        # inside the HTML structure.
        if node.name is None:

            content = str(node)

            # Skip empty or whitespace-only nodes to avoid
            # generating unnecessary DOCX runs.
            if not content.strip():
                return

            run = paragraph.add_run(content)

            # Apply inherited formatting to the run
            if formatting.get("bold"):
                run.bold = True

            if formatting.get("italic"):
                run.italic = True

            if formatting.get("underline"):
                run.underline = True

            if formatting.get("strike"):
                run.font.strike = True

            # Apply text color if present
            rgb = formatting.get("color")

            if isinstance(rgb, str) and len(rgb) == 6:
                run.font.color.rgb = RGBColor.from_string(rgb)

        else:

            # ----------------------------------------
            # ELEMENT NODE
            # ----------------------------------------
            # When encountering an HTML element (<b>, <span>, etc.)
            # we update the active formatting before processing children.
            # Child nodes inherit this updated formatting.
            new_format = formatting.copy()

            if node.name == "a":

                url = node.get("href")

                # Create empty hyperlink first
                hyperlink_text = ""

                for child in node.children:
                    hyperlink_text += str(child)

                add_hyperlink(paragraph, url, hyperlink_text)

                return

            if node.name in ["strong", "b"]:
                new_format["bold"] = True

            if node.name in ["em", "i"]:
                new_format["italic"] = True

            if node.name == "u":
                new_format["underline"] = True

            if node.name == "strike":
                new_format["strike"] = True

            if node.name in ["span", "font"]:
                match = re.search(r"#([0-9A-Fa-f]{6})", str(node))
                if match:
                    new_format["color"] = match.group(1)

            # Recursively process child nodes so formatting cascades
            for child in node.children:
                process_node(child, paragraph, new_format)

    # ----------------------------------------
    # Build DOCX paragraphs
    # ----------------------------------------
    # Paragraphs in the template engine are separated by double line breaks.
    paragraphs = text.split("\n\n")

    for p in paragraphs:

        paragraph = doc.add_paragraph()

        # Parse paragraph HTML so formatting can be processed node-by-node
        soup = BeautifulSoup(p, "html.parser")

        # Each top-level node inside the paragraph is processed
        for child in soup.children:
            process_node(child, paragraph)

    buffer = io.BytesIO()
    doc.save(buffer)

    return buffer.getvalue()


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Helper function to convert a Word docx to pdf bytes
    """

    with tempfile.TemporaryDirectory() as tmpdir:

        docx_path = f"{tmpdir}/file.docx"
        pdf_path = f"{tmpdir}/file.pdf"

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        convert(docx_path, pdf_path)

        with open(pdf_path, "rb") as f:
            return f.read()


def normalize_key(value: str) -> str:
    """
    Normalize strings for reliable key comparisons.

    The function removes whitespace, punctuation and converts Danish
    characters so template keys can be matched regardless of formatting
    differences between Excel, API data and placeholders.

    Args:
        value (str): Input string.

    Returns:
        str: Normalized key used for comparisons.
    """

    return (
        value.strip()
        .lower()
        .replace(" ", "")
        .replace(".", "")
        .replace("ø", "oe")
        .replace("å", "aa")
        .replace("æ", "ae")
        .replace("?", "")
        .replace("-", "")
        .replace("_", "")
    )


def replace_placeholders(text: str, data: dict) -> str:
    """
    Replace placeholders in the form {key} with values from the data dictionary.

    The function also cleans malformed placeholders that may occur when Excel
    formatting wraps placeholders in HTML tags. Replaced values are wrapped in
    a blue span so inserted data is visually distinguishable in the output.

    Args:
        text (str): Letter text containing placeholders.
        data (dict): Data used to resolve placeholder values.

    Returns:
        str: Text with placeholders replaced.
    """

    # ----------------------------------------
    # Fix malformed placeholders
    # ----------------------------------------
    # Excel rich text formatting can sometimes produce placeholders like:
    # {<span>barnets_fornavn</span>}
    # This regex removes the HTML tags inside the placeholder.
    text = re.sub(
        r"\{<[^>]+>(.*?)</[^>]+>\}",
        r"{\1}",
        text
    )

    def repl(match):

        # Extract placeholder key and clean invisible characters
        key = match.group(1).replace("\u200b", "").strip()

        value = data.get(key)

        # If no value exists, keep the original placeholder
        if value is None:
            return match.group(0)

        # Wrap replacements in a blue color span so inserted values are visually distinguishable in the final document.
        return f'<span style="color:#0F9ED5">{value}</span>'

    # Replace all placeholders of the form {key}
    return re.sub(r"\{([^{}]+)\}", repl, text)
